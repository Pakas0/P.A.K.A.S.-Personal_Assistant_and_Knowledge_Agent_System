from docx import Document

def build_docx(title: str, sections: list[dict], output_path: str) -> str:
    """
    sections: list of {"heading": str, "content": str} — struktur sederhana,
    heading jadi Heading level 1/2, content jadi paragraf biasa.
    Return output_path setelah file berhasil ditulis.
    """
    doc = Document()
    doc.add_heading(title, level=0)
    for sec in sections:
        if sec.get("heading"):
            doc.add_heading(sec["heading"], level=1)
        doc.add_paragraph(sec.get("content", ""))
    doc.save(output_path)
    return output_path
